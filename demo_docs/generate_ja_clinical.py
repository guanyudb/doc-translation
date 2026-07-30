"""Generate synthetic Japanese clinical DOCX files for demoing the
doc-translation pipeline end to end.

Each file targets ~10 pages of realistic clinical content to properly
exercise the OOXML in-place translator (paragraphs, tables, running
header/footer, native tables, embedded raster figures).

All content is 100% synthetic:
  * Drug name "DBX-101" and study codes are invented.
  * Company "Databricks Pharma KK" is fictional.
  * Patient IDs, lab values, and adverse-event counts are made up.
  * Document structure follows publicly-published regulator templates:
      - ICH E6 (R2) — Good Clinical Practice (protocol synopsis structure)
      - ICH E2C (R2) — Periodic Safety Update (PSUR structure)
      - ICH M4E — Common Technical Document (CTD Module 2.7.3 structure)
    These ICH guidelines are public at database.ich.org.
  * Japanese clinical terminology is generic (公知の医学用語) — no
    customer-derived vocabulary.

Elements exercised (per the translator design):
  * <w:p> body paragraphs               — translated
  * <w:tbl>/<w:tc> tables               — translated
  * word/header*.xml + footer*.xml      — translated
  * embedded raster figures (<w:drawing>) — intentionally skipped
    (mimics real clinical charts pasted from Excel/GraphPad/SAS)
"""
from __future__ import annotations
from pathlib import Path
import io

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib import font_manager
import numpy as np

from docx import Document
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Inches, Pt, RGBColor


OUT_DIR = Path(__file__).parent / "ja_clinical"
OUT_DIR.mkdir(parents=True, exist_ok=True)


# ---------------------------------------------------------------------------
# Chart helpers — matplotlib PNGs embedded as figures
# ---------------------------------------------------------------------------

def _mpl_font():
    for name in ("Hiragino Sans", "Yu Gothic", "Noto Sans CJK JP", "IPAGothic", "MS Gothic"):
        for f in font_manager.fontManager.ttflist:
            if f.name == name:
                return name
    return None

_JP_FONT = _mpl_font()
if _JP_FONT:
    plt.rcParams["font.family"] = _JP_FONT

def _png_bytes(fig) -> bytes:
    buf = io.BytesIO(); fig.savefig(buf, format="png", dpi=140); plt.close(fig)
    return buf.getvalue()


def _pk_profile_png() -> bytes:
    t = np.array([0, 0.5, 1, 1.5, 2, 3, 4, 6, 8, 12, 16, 24])
    dose_100 = np.array([0, 45, 120, 180, 210, 195, 165, 120, 85, 45, 22, 8])
    dose_200 = np.array([0, 88, 235, 355, 420, 385, 320, 240, 175, 90, 45, 18])
    fig, ax = plt.subplots(figsize=(6.5, 3.6))
    ax.plot(t, dose_100, marker="o", label="100 mg")
    ax.plot(t, dose_200, marker="s", label="200 mg")
    ax.set_xlabel("投与後時間 (h)" if _JP_FONT else "Time after dose (h)")
    ax.set_ylabel("血漿中濃度 (ng/mL)" if _JP_FONT else "Plasma concentration (ng/mL)")
    ax.set_title("DBX-101 単回投与時 薬物動態プロファイル" if _JP_FONT else "DBX-101 single-dose PK profile")
    ax.legend(); ax.grid(True, alpha=0.3); fig.tight_layout()
    return _png_bytes(fig)


def _teae_by_dose_png() -> bytes:
    grades = ["Grade 1", "Grade 2", "Grade 3", "Grade 4", "Grade 5"]
    n_100 = [22, 15, 6, 1, 0]; n_200 = [28, 20, 10, 3, 1]; n_ctl = [20, 12, 4, 1, 0]
    x = np.arange(len(grades)); w = 0.27
    fig, ax = plt.subplots(figsize=(6.5, 3.6))
    ax.bar(x - w, n_100, w, label="100 mg")
    ax.bar(x,      n_200, w, label="200 mg")
    ax.bar(x + w,  n_ctl, w, label="対照群" if _JP_FONT else "Control")
    ax.set_xticks(x); ax.set_xticklabels(grades)
    ax.set_ylabel("症例数" if _JP_FONT else "n patients")
    ax.set_title("投与群別 CTCAE Grade 別 有害事象発現例数" if _JP_FONT
                 else "TEAE grade by dose group")
    ax.legend(); ax.grid(True, axis="y", alpha=0.3); fig.tight_layout()
    return _png_bytes(fig)


def _ae_bar_png() -> bytes:
    socs = ["消化器", "神経系", "皮膚", "肝胆道", "血液系", "全身状態"] if _JP_FONT \
           else ["GI", "Nervous", "Skin", "Hepatobiliary", "Blood", "General"]
    n_100 = [18, 12, 9, 4, 3, 8]; n_200 = [26, 17, 13, 7, 5, 11]
    x = np.arange(len(socs)); w = 0.38
    fig, ax = plt.subplots(figsize=(6.5, 3.6))
    ax.bar(x - w/2, n_100, w, label="100 mg (n=60)")
    ax.bar(x + w/2, n_200, w, label="200 mg (n=60)")
    ax.set_xticks(x); ax.set_xticklabels(socs, rotation=15)
    ax.set_ylabel("発現例数" if _JP_FONT else "Number of events")
    ax.set_title("器官別大分類 (SOC) 別 有害事象頻度" if _JP_FONT else "AE frequency by SOC")
    ax.legend(); ax.grid(True, axis="y", alpha=0.3); fig.tight_layout()
    return _png_bytes(fig)


def _safety_trend_png() -> bytes:
    periods = ["2024下期", "2025上期", "2025下期", "2026上期"] if _JP_FONT \
              else ["H2 2024", "H1 2025", "H2 2025", "H1 2026"]
    exposure = [420, 780, 1120, 1560]
    saes = [6, 11, 16, 21]
    fig, ax1 = plt.subplots(figsize=(6.5, 3.6))
    ax2 = ax1.twinx()
    ax1.bar(periods, exposure, color="steelblue", alpha=0.7,
            label="患者年" if _JP_FONT else "Patient-years")
    ax2.plot(periods, saes, color="crimson", marker="o", linewidth=2,
             label="重篤AE件数" if _JP_FONT else "Serious AEs")
    ax1.set_ylabel("累積曝露量 (患者年)" if _JP_FONT else "Cumulative exposure (patient-years)",
                   color="steelblue")
    ax2.set_ylabel("重篤有害事象件数" if _JP_FONT else "Serious AEs (count)", color="crimson")
    ax1.set_title("報告期間別 曝露量および重篤有害事象件数" if _JP_FONT
                  else "Cumulative exposure & serious AEs by reporting period")
    ax1.grid(True, axis="y", alpha=0.3); fig.tight_layout()
    return _png_bytes(fig)


def _forest_png() -> bytes:
    subgroups = ["全体", "年齢 65歳未満", "年齢 65歳以上", "男性", "女性", "軽症", "重症",
                 "PS 0", "PS 1", "前治療1レジメン", "前治療2レジメン以上"] \
        if _JP_FONT else ["Overall", "Age<65", "Age>=65", "Male", "Female", "Mild", "Severe",
                          "PS 0", "PS 1", "Prior 1L", "Prior 2L+"]
    hr = np.array([0.68, 0.62, 0.75, 0.71, 0.65, 0.58, 0.79, 0.61, 0.74, 0.55, 0.78])
    lo = np.array([0.55, 0.47, 0.58, 0.55, 0.49, 0.42, 0.60, 0.44, 0.58, 0.39, 0.61])
    hi = np.array([0.84, 0.83, 0.98, 0.92, 0.87, 0.80, 1.04, 0.85, 0.95, 0.78, 0.99])
    y  = np.arange(len(subgroups))[::-1]
    fig, ax = plt.subplots(figsize=(6.5, 4.2))
    ax.errorbar(hr, y, xerr=[hr-lo, hi-hr], fmt="s", capsize=4, color="tab:blue")
    ax.axvline(1.0, color="grey", linestyle="--", linewidth=0.8)
    ax.set_yticks(y); ax.set_yticklabels(subgroups)
    ax.set_xlabel("ハザード比 (95% CI)" if _JP_FONT else "Hazard ratio (95% CI)")
    ax.set_title("部分集団別 主要評価項目 (無増悪生存期間) ハザード比" if _JP_FONT
                 else "Primary endpoint HR by subgroup (PFS)")
    ax.set_xlim(0.3, 1.3); ax.grid(True, axis="x", alpha=0.3); fig.tight_layout()
    return _png_bytes(fig)


def _km_curve_png() -> bytes:
    t = np.linspace(0, 24, 100)
    s_dbx200 = np.exp(-t / 15.0)
    s_dbx100 = np.exp(-t / 12.0)
    s_ctl    = np.exp(-t / 9.4)
    fig, ax = plt.subplots(figsize=(6.5, 3.6))
    ax.step(t, s_dbx200, where="post", label="DBX-101 200 mg", linewidth=2)
    ax.step(t, s_dbx100, where="post", label="DBX-101 100 mg", linewidth=2)
    ax.step(t, s_ctl,    where="post", label="対照群" if _JP_FONT else "Control", linewidth=2)
    ax.axhline(0.5, color="grey", linestyle=":", linewidth=0.8)
    ax.set_xlabel("観察期間 (月)" if _JP_FONT else "Months")
    ax.set_ylabel("無増悪生存確率" if _JP_FONT else "PFS probability")
    ax.set_title("Kaplan-Meier 無増悪生存曲線 (投与群別)" if _JP_FONT
                 else "Kaplan-Meier PFS curve by treatment arm")
    ax.set_ylim(0, 1.02); ax.legend(); ax.grid(True, alpha=0.3); fig.tight_layout()
    return _png_bytes(fig)


def _waterfall_png() -> bytes:
    rng = np.random.default_rng(42)
    n = 55
    change = np.concatenate([
        rng.normal(-45, 20, n // 2),
        rng.normal(10,  25, n - n // 2),
    ])
    change = np.sort(change)
    colors = ["tab:green" if c < -30 else ("tab:orange" if c < 20 else "tab:red") for c in change]
    fig, ax = plt.subplots(figsize=(6.5, 3.6))
    ax.bar(np.arange(n), change, color=colors)
    ax.axhline(-30, color="grey", linestyle="--", linewidth=0.8)
    ax.axhline(20,  color="grey", linestyle="--", linewidth=0.8)
    ax.set_xlabel("被験者 (最大変化率順)" if _JP_FONT else "Subjects (sorted by best change)")
    ax.set_ylabel("腫瘍径合計 最大変化率 (%)" if _JP_FONT else "Best % change in sum of diameters")
    ax.set_title("腫瘍縮小効果 ウォーターフォールプロット (DBX-101 200 mg 群)" if _JP_FONT
                 else "Tumor response waterfall (DBX-101 200 mg)")
    ax.grid(True, axis="y", alpha=0.3); fig.tight_layout()
    return _png_bytes(fig)


# ---------------------------------------------------------------------------
# Header / footer + table helpers
# ---------------------------------------------------------------------------

def _add_header_footer(doc: Document, header_text: str, footer_text: str) -> None:
    section = doc.sections[0]
    hdr = section.header.paragraphs[0]
    hdr.text = header_text
    hdr.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in hdr.runs:
        run.font.size = Pt(9); run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    ftr = section.footer.paragraphs[0]
    ftr.text = footer_text + "  |  ページ "
    ftr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in ftr.runs:
        run.font.size = Pt(9); run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run = ftr.add_run(); run.font.size = Pt(9)
    fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r"); t = OxmlElement("w:t"); t.text = "1"; r.append(t); fld.append(r)
    run._r.append(fld)


def _shade_row(row, hex_color: str) -> None:
    for cell in row.cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hex_color)
        tc_pr.append(shd)


def _fill_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            for r in p.runs: r.bold = True
    _shade_row(tbl.rows[0], "D9E2F3")
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            tbl.rows[r_idx].cells[c_idx].text = str(val)
    for row in tbl.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                for r in p.runs: r.font.size = Pt(10)


def _h(doc, text, level=1): doc.add_heading(text, level=level)

def _p(doc, text):
    para = doc.add_paragraph(text)
    for r in para.runs: r.font.size = Pt(11)

def _fig(doc, png_bytes: bytes, caption: str) -> None:
    doc.add_picture(io.BytesIO(png_bytes), width=Inches(6.0))
    para = doc.add_paragraph(caption); para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in para.runs:
        r.italic = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


# ---------------------------------------------------------------------------
# Doc 1 — Clinical Trial Protocol Synopsis (JA, ~10 pages)
# Structure: ICH E6 (R2) — Good Clinical Practice
# ---------------------------------------------------------------------------

def build_protocol_synopsis() -> Path:
    doc = Document()
    _add_header_footer(
        doc,
        header_text="治験実施計画書  概要  |  DBX101-JP-002  |  版 2.0 (2026年3月15日)",
        footer_text="Databricks Pharma KK  秘密扱い",
    )

    _h(doc, "治験実施計画書 概要", 0)
    _p(doc, "本書は、DBX-101 (開発コード) を対象とする第II相無作為化二重盲検比較試験の治験実施計画書の概要である。"
             "本試験は、日本国内10施設で実施される多施設共同試験である。")
    _p(doc, "本概要は、ICH E6 (R2) Good Clinical Practice ガイドラインおよび"
             "医薬品、医療機器等の品質、有効性及び安全性の確保等に関する法律 (医薬品医療機器等法)"
             "に基づいて作成されている。")

    # ---- 1. Background & Rationale ----
    _h(doc, "1. 治験の背景および根拠", 1)
    _h(doc, "1.1 疾患背景", 2)
    _p(doc, "本疾患領域における現行の標準治療は、5年生存率が約35%に留まっており、新規治療選択肢に対する強い医療ニーズが存在する。"
             "特に標準的な一次治療に不応となった患者では、二次治療以降の選択肢が限られており、"
             "予後改善に寄与する新規作用機序の薬剤が求められている。")
    _p(doc, "近年、当該経路を標的とする分子標的治療の可能性が示唆されており、複数の前臨床研究において"
             "腫瘍増殖抑制作用が報告されている。臨床開発においては、単剤療法および併用療法の両方向で開発が進められている。")

    _h(doc, "1.2 DBX-101 の作用機序", 2)
    _p(doc, "DBX-101 は新規経路を選択的に阻害する小分子化合物であり、"
             "生化学的アッセイにおいて標的酵素に対するIC50値は 8.4 nMであることが確認されている。"
             "細胞株実験では、対象疾患由来細胞株において用量依存的な増殖抑制作用が観察されている。")
    _p(doc, "動物モデルにおいては、経口投与により腫瘍縮小効果が確認されており、"
             "薬物動態学的にはヒト用量に換算して 200 mg/day 前後で有効血中濃度に到達すると予測されている。"
             "第I相試験 (DBX101-JP-001) では、100 mg および 200 mg 用量で忍容性が確認された。")

    _h(doc, "1.3 本試験実施の根拠", 2)
    _p(doc, "第I相試験の結果を踏まえ、本試験では、標準治療に不応または不耐となった患者を対象に、"
             "DBX-101の有効性と安全性を無作為化比較試験により評価する。"
             "得られた結果は、後続の第III相試験の症例数設計、投与量選択、およびエンドポイント設定に用いられる。")

    # ---- 2. Objectives ----
    _h(doc, "2. 治験の目的", 1)
    _h(doc, "2.1 主要目的", 2)
    _p(doc, "DBX-101 (100 mgおよび200 mgの2用量) の、標準治療との比較における無増悪生存期間 (PFS) の延長効果を評価する。")

    _h(doc, "2.2 副次目的", 2)
    _p(doc, "以下の副次目的を評価する:")
    _p(doc, "① 全生存期間 (OS) の延長効果")
    _p(doc, "② 客観的奏効率 (ORR) および奏効期間 (DOR) の改善")
    _p(doc, "③ 病勢制御率 (DCR) の改善")
    _p(doc, "④ 患者報告アウトカム (PRO) に対する影響")
    _p(doc, "⑤ 安全性プロファイルの確認 (CTCAE Grade別発現頻度)")
    _p(doc, "⑥ 薬物動態パラメータ (Cmax、AUC、t1/2) の推定")

    _h(doc, "2.3 探索的目的", 2)
    _p(doc, "腫瘍組織および血漿検体を用いたバイオマーカー解析を実施し、"
             "DBX-101の効果予測因子および耐性因子の同定を試みる。"
             "また、循環腫瘍DNA (ctDNA) の変化と臨床効果の相関を探索する。")

    # ---- 3. Design ----
    _h(doc, "3. 治験デザイン", 1)
    _p(doc, "本試験は、多施設共同、無作為化、二重盲検、対照群設定の第II相試験である。"
             "適格患者は、1:1:1の比率でDBX-101 100 mg群、DBX-101 200 mg群、または対照群 (標準治療) に無作為に割付けられる。")
    _p(doc, "無作為化は施設、ECOGパフォーマンスステータス (0 vs 1)、および前治療レジメン数 (1 vs 2以上)"
             "を層別因子として実施する。二重盲検性を維持するため、対照群には対応するプラセボが投与される。")

    _fill_table(doc,
        ["項目", "内容"],
        [
            ["試験相",       "第II相"],
            ["試験タイプ",   "介入試験"],
            ["デザイン",     "多施設共同、無作為化、二重盲検、対照群比較"],
            ["割付比率",     "1 : 1 : 1"],
            ["層別因子",     "施設、ECOG PS (0 vs 1)、前治療数 (1 vs 2+)"],
            ["予定登録数",   "180例 (60例 × 3群)"],
            ["予定試験期間", "登録開始から3年間 (登録期間2年 + 追跡1年)"],
            ["中間解析",     "PFSイベント発生70%時点で1回実施"],
        ])
    doc.add_paragraph()

    # ---- 4. Study Population ----
    _h(doc, "4. 対象患者", 1)
    _h(doc, "4.1 選択基準", 2)
    _p(doc, "以下の全てを満たす患者を組み入れる:")
    _p(doc, "① 同意取得時に20歳以上85歳以下の成人男女")
    _p(doc, "② 組織学的または細胞学的に対象疾患と確定診断されていること")
    _p(doc, "③ RECIST v1.1 基準による測定可能病変を1つ以上有すること")
    _p(doc, "④ ECOG パフォーマンスステータスが 0 または 1 であること")
    _p(doc, "⑤ 標準治療に不応または不耐となった患者 (前治療 1〜4 レジメン)")
    _p(doc, "⑥ 予測生存期間が3か月以上であること")
    _p(doc, "⑦ 十分な臓器機能を有すること (詳細は治験実施計画書本文に記載)")
    _p(doc, "⑧ 経口投与が可能であること")
    _p(doc, "⑨ 妊娠可能な女性および女性パートナーを有する男性は避妊に同意すること")
    _p(doc, "⑩ 本人による文書同意が取得できていること")

    _h(doc, "4.2 除外基準", 2)
    _p(doc, "以下のいずれかに該当する患者は除外する:")
    _p(doc, "① 妊娠中または授乳中の女性")
    _p(doc, "② 活動性の二次悪性腫瘍を有する患者 (適切に治療された皮膚基底細胞癌等を除く)")
    _p(doc, "③ 中枢神経系転移を有し、症候性かつ治療されていない患者")
    _p(doc, "④ コントロール不良な心血管疾患 (登録前6か月以内の心筋梗塞、不安定狭心症等) を有する患者")
    _p(doc, "⑤ 臨床的に問題となる不整脈 (QTc延長を含む) を有する患者")
    _p(doc, "⑥ 活動性の感染症 (HIV、HBV、HCV、活動性結核) を有する患者")
    _p(doc, "⑦ 試験薬剤の成分に対する既知の過敏症を有する患者")
    _p(doc, "⑧ 前治療の毒性が Grade 1 以下に回復していない患者 (脱毛等を除く)")
    _p(doc, "⑨ 登録前 14 日以内に他の治験薬による治療を受けた患者")
    _p(doc, "⑩ 治験実施計画書の遵守が困難と判断される精神疾患を有する患者")
    _p(doc, "⑪ その他、治験責任医師が不適当と判断した患者")

    # ---- 5. Treatments ----
    _h(doc, "5. 用法・用量および投与スケジュール", 1)
    _fill_table(doc,
        ["コホート", "用量", "投与経路", "投与頻度", "投与期間"],
        [
            ["低用量群 (100 mg)", "100 mg", "経口",     "1日1回 (朝食後)", "疾患進行または投与中止まで"],
            ["高用量群 (200 mg)", "200 mg", "経口",     "1日1回 (朝食後)", "疾患進行または投与中止まで"],
            ["対照群",           "標準治療", "指定なし", "各薬剤の添付文書に従う", "疾患進行または投与中止まで"],
        ])
    doc.add_paragraph()

    _h(doc, "5.1 用量調整基準", 2)
    _p(doc, "有害事象の重症度に応じて、以下の基準に従い用量調整または休薬を行う。"
             "Grade 3 以上の非血液毒性、または Grade 4 の血液毒性が発現した場合は、"
             "Grade 1 以下に回復するまで休薬する。休薬期間が 21 日を超える場合は、投与を永続的に中止する。")
    _fill_table(doc,
        ["有害事象", "初回発現", "2回目発現", "3回目発現"],
        [
            ["Grade 3 非血液毒性", "1レベル減量 (100 mg)", "1レベル減量 (50 mg)", "投与中止"],
            ["Grade 4 血液毒性",   "1レベル減量",           "1レベル減量",           "投与中止"],
            ["QTc > 500 ms",     "休薬 → 500 ms未満で1レベル減量再開", "投与中止", "―"],
            ["肝機能異常 (AST/ALT > 5×ULN)", "休薬 → 3×ULN未満で1レベル減量再開", "投与中止", "―"],
        ])
    doc.add_paragraph()

    _h(doc, "5.2 併用療法および禁止療法", 2)
    _p(doc, "併用可能: 支持療法薬 (制吐剤、下痢止め、鎮痛剤等)、基礎疾患に対する既存治療の継続。"
             "禁止: 他の抗腫瘍薬、免疫抑制薬 (プロトコル指定を除く)、強力な CYP3A4 阻害薬・誘導薬、"
             "および試験期間中の生ワクチン接種。")

    # ---- 6. Endpoints ----
    _h(doc, "6. 評価項目", 1)
    _p(doc, "主要評価項目: 独立中央判定 (BICR) による RECIST v1.1 基準に基づく無増悪生存期間 (PFS)。"
             "PFS は無作為化割付日から、疾患進行または全ての原因による死亡のいずれか早い方の日までの期間と定義する。")
    _p(doc, "副次評価項目:")
    _p(doc, "① 全生存期間 (OS): 無作為化割付日から死亡日までの期間")
    _p(doc, "② 客観的奏効率 (ORR): RECIST v1.1 における CR + PR の割合")
    _p(doc, "③ 奏効期間 (DOR): 初回奏効確認日から進行/死亡までの期間")
    _p(doc, "④ 病勢制御率 (DCR): CR + PR + SD の割合")
    _p(doc, "⑤ 有害事象の発現頻度および重症度 (CTCAE v5.0)")
    _p(doc, "⑥ 患者報告アウトカム: EORTC QLQ-C30 スコアの経時変化")
    _p(doc, "⑦ 薬物動態パラメータ (Cmax, AUC0-24, t1/2, Ctrough)")

    # ---- 7. Assessments schedule ----
    _h(doc, "7. 検査・評価スケジュール", 1)
    _p(doc, "各来院時に実施する検査・評価の概要を以下に示す。詳細は治験実施計画書本文の"
             "「検査・評価スケジュール表 (Schedule of Assessments)」を参照する。")
    _fill_table(doc,
        ["評価項目", "スクリーニング (-28日〜-1日)", "サイクル1", "サイクル2以降", "追跡調査"],
        [
            ["同意取得",           "○", "―", "―", "―"],
            ["病歴・身体診察",     "○", "○", "○ (各サイクル初日)", "○"],
            ["バイタルサイン",     "○", "○", "○", "○"],
            ["ECOG PS",           "○", "○", "○", "○"],
            ["血液学的検査",       "○", "○", "○ (2週毎)", "○"],
            ["生化学的検査",       "○", "○", "○ (2週毎)", "○"],
            ["尿検査",            "○", "○", "○ (4週毎)", "―"],
            ["心電図 (12誘導)",   "○", "○", "○ (4週毎)", "○"],
            ["画像評価 (CT/MRI)", "○", "―", "○ (8週毎)", "○"],
            ["QLQ-C30",           "○", "○", "○ (各サイクル初日)", "○"],
            ["PK 採血",           "―", "○", "○ (Cycle 2, 4)", "―"],
            ["ctDNA 採血",        "○", "○", "○ (Cycle 3, 6)", "○"],
            ["有害事象評価",      "―", "継続", "継続", "継続 (30日間)"],
        ])
    doc.add_paragraph()

    # ---- 8. PK ----
    _h(doc, "8. 薬物動態評価", 1)
    _p(doc, "サイクル1初日 (単回投与) およびサイクル1第15日 (反復投与定常状態) において、"
             "投与前および投与後 0.5、1、2、4、8、12、24 時間の合計8点で PK 採血を実施する。"
             "サイクル2および4の初日には、投与前 (Ctrough) の1点のみ採血する。"
             "血漿中 DBX-101 濃度は、バリデート済みの LC-MS/MS 法により測定する。")
    _fig(doc, _pk_profile_png(),
         "図1. DBX-101 単回経口投与時の血漿中濃度推移 (想定値、模擬データ)")
    _p(doc, "予備解析の結果、DBX-101 の薬物動態はほぼ線形であり、"
             "用量比例的に Cmax および AUC が増加すると予測される。"
             "反復投与時の蓄積係数は約 1.4 と推定される。")

    # ---- 9. Safety expectations ----
    _h(doc, "9. 安全性評価", 1)
    _p(doc, "全ての有害事象は、CTCAE v5.0 に基づいて Grade 判定を行う。"
             "第I相試験の結果から予測される主な有害事象は、消化器症状 (悪心、下痢、食欲低下)、"
             "全身状態 (倦怠感、疲労)、および皮膚関連 (発疹、掻痒感) である。"
             "予測される重篤有害事象としては、Grade 3 以上の肝機能異常、および Grade 3 以上の好中球減少が挙げられる。")
    _fig(doc, _teae_by_dose_png(),
         "図2. 投与群別 CTCAE Grade 別 有害事象発現例数 (想定値、模擬データ)")

    # ---- 10. Statistical analysis ----
    _h(doc, "10. 統計解析計画", 1)
    _h(doc, "10.1 症例数設計", 2)
    _p(doc, "対照群の PFS 中央値を 6.5 か月、DBX-101 高用量群の PFS 中央値を 10.4 か月と仮定した場合、"
             "片側有意水準 0.025、検出力 80%、追跡打切り率 15% を仮定し、"
             "必要症例数は各群 60 例、合計 180 例と算出された。")

    _h(doc, "10.2 主要解析", 2)
    _p(doc, "主要解析集団は無作為化された全患者からなる ITT 集団とする。"
             "群間の PFS 比較は、層別因子で層別化した Log-rank 検定を用いて実施する。"
             "ハザード比および 95% 信頼区間は、層別化 Cox 比例ハザードモデルにより推定する。")

    _h(doc, "10.3 中間解析", 2)
    _p(doc, "PFS イベント発生数が計画総イベント数の 70% に到達した時点で、"
             "独立データモニタリング委員会 (IDMC) による中間解析を実施する。"
             "無効性中止基準として、条件付検出力が 10% 未満となった場合は試験中止を勧告する。"
             "α消費関数は O'Brien-Fleming 型を採用する。")

    _h(doc, "10.4 感度解析および部分集団解析", 2)
    _p(doc, "主要評価項目については、per-protocol 集団を対象とした感度解析、"
             "および事前に規定された部分集団 (年齢、性別、ECOG PS、前治療数等)"
             "における Cox モデルによる部分集団解析を実施する。")

    # ---- 11. DMC / Ethics ----
    _h(doc, "11. データモニタリング委員会", 1)
    _p(doc, "本試験には独立データモニタリング委員会 (IDMC) を設置する。"
             "IDMC は、当該領域に精通した治験責任医師 3 名 (うち 1 名は生物統計家) から構成され、"
             "登録開始後 6 か月毎、および中間解析時点において、安全性および有効性データをレビューする。")

    _h(doc, "12. 倫理的および規制上の考慮事項", 1)
    _p(doc, "本試験は、ヘルシンキ宣言、ICH E6 (R2) Good Clinical Practice、"
             "および医薬品医療機器等法を遵守して実施される。"
             "試験開始前に、各実施医療機関の治験審査委員会 (IRB) の承認、"
             "および医薬品医療機器総合機構 (PMDA) への治験届出を完了する。")
    _p(doc, "全ての被験者は、試験参加前に治験の目的、方法、予想される利益および不利益について、"
             "治験責任医師から十分な説明を受け、自由意思による文書同意を提供する。"
             "被験者は、いかなる不利益を被ることなく、いつでも同意を撤回する権利を有する。")

    _h(doc, "13. 試験組織および連絡先", 1)
    _p(doc, "治験調整医師: 本試験の医学的な統括を担当する。"
             "生物統計責任者: 統計解析計画書の作成および統計解析を担当する。"
             "データマネジメント責任者: EDC の管理およびデータクリーニングを担当する。"
             "モニタリング担当: CRA による定期的なモニタリング訪問を実施する。")

    out = OUT_DIR / "01_protocol_synopsis_ja.docx"
    doc.save(out)
    return out


# ---------------------------------------------------------------------------
# Doc 2 — Periodic Safety Update Report / PSUR (JA, ~10 pages)
# Structure: ICH E2C (R2)
# ---------------------------------------------------------------------------

def build_ae_summary() -> Path:
    doc = Document()
    _add_header_footer(
        doc,
        header_text="定期的安全性最新報告書 (PSUR) 概要  |  DBX-101  |  報告期間 2025年10月1日 〜 2026年3月31日",
        footer_text="Databricks Pharma KK  安全性管理部",
    )

    _h(doc, "定期的安全性最新報告書 (PSUR)", 0)
    _p(doc, "本書は、DBX-101 (開発コード) について、報告期間中に収集された安全性情報を要約したものである。"
             "本要約は、ICH E2C (R2) Periodic Benefit-Risk Evaluation Report ガイドラインに準拠して作成されている。")

    _h(doc, "1. はじめに", 1)
    _p(doc, "本報告書の対象期間は 2025 年 10 月 1 日から 2026 年 3 月 31 日までの 6 か月間である。"
             "本剤の国内販売承認取得日は 2024 年 4 月 15 日であり、国際誕生日 (IBD) を基準とした"
             "累積 PSUR 番号は本報告書で第 4 号となる。")
    _p(doc, "本報告書に記載する安全性情報は、"
             "自発報告、市販後調査、実施中の臨床試験、および公表文献レビューから収集されたものである。"
             "全てのデータは、報告期間終了日である 2026 年 3 月 31 日のデータロック時点でクリーニング済みである。")

    _h(doc, "2. 医薬品の販売承認状況", 1)
    _p(doc, "報告期間終了時点の DBX-101 の世界における販売承認状況を以下に示す。")
    _fill_table(doc,
        ["国・地域", "承認日", "適応症", "販売状況", "備考"],
        [
            ["日本",     "2024-04-15", "対象疾患",     "販売中",     "IBD"],
            ["米国",     "2024-08-22", "対象疾患",     "販売中",     "―"],
            ["EU",      "2024-11-10", "対象疾患",     "販売中",     "―"],
            ["韓国",     "2025-06-03", "対象疾患",     "販売中",     "―"],
            ["台湾",     "2025-09-18", "対象疾患",     "販売中",     "―"],
            ["中国",     "審査中",     "―",           "承認申請中", "PMDA相談完了"],
        ])
    doc.add_paragraph()

    _h(doc, "3. 安全上の理由による措置", 1)
    _p(doc, "報告期間中、規制上の措置として、EU 規制当局からの要請により添付文書の 4.4 項"
             "「使用上の注意」に、Grade 3 以上の肝機能異常発現時の対応に関する記載を追加した。"
             "その他、リコール、販売停止、および流通制限に至る措置は発生していない。")

    _h(doc, "4. 添付文書 (RSI) の変更", 1)
    _p(doc, "報告期間中、以下の項目について添付文書の改訂を実施した:"
             "(1) 4.4 項「特別な警告および使用上の注意」に肝機能異常の管理に関する記載追加。"
             "(2) 4.8 項「副作用」の頻度データを、報告期間中の累積データに基づいて更新。"
             "(3) 5.3 項「非臨床試験データ」に長期毒性試験の追加データを反映。")

    _h(doc, "5. 推定累積曝露量", 1)
    _p(doc, "報告期間中および累積での DBX-101 の推定曝露量を以下に示す。"
             "曝露量は、販売実績データおよび推定平均投与量 (150 mg/day) を基に算出した。")
    _fill_table(doc,
        ["区分", "本報告期間 (2025-10-01 〜 2026-03-31)", "累積 (承認以降)"],
        [
            ["投与患者数 (推定)",    "3,240例",    "12,860例"],
            ["患者年 (推定)",       "1,560患者年", "3,880患者年"],
            ["日本国内",           "1,120例",    "5,240例"],
            ["日本国外",           "2,120例",    "7,620例"],
            ["臨床試験参加者",      "180例",      "620例"],
        ])
    doc.add_paragraph()
    _fig(doc, _safety_trend_png(),
         "図1. 報告期間別 累積曝露量および重篤有害事象件数の推移 (想定値、模擬データ)")

    _h(doc, "6. 報告期間中の総合安全性評価", 1)
    _p(doc, "報告期間中に自発報告として収集された有害事象は 428 件 (国内 178 件、海外 250 件) であり、"
             "そのうち重篤例は 62 件 (14.5%)、死亡例は 8 件であった。"
             "死亡例のうち、本剤との因果関係が否定できない症例は 2 例であった。")
    _p(doc, "全体として、DBX-101 の安全性プロファイルは前回報告書から大きな変化はなく、"
             "新たな安全性シグナルは検出されていない。"
             "最も頻度の高い有害事象は消化器症状 (悪心、下痢、食欲低下) であり、大半は Grade 1 〜 2 であった。")

    _h(doc, "7. 器官別大分類 (SOC) 別 有害事象一覧", 1)
    _fill_table(doc,
        ["器官別大分類 (SOC)", "報告例数", "重篤例数", "累積例数", "累積重篤例数"],
        [
            ["胃腸障害",                        "142 (33.2%)", "8",  "1,845", "62"],
            ["神経系障害",                      "78 (18.2%)",  "3",  "980",   "22"],
            ["皮膚および皮下組織障害",           "63 (14.7%)",  "1",  "820",   "12"],
            ["肝胆道系障害",                    "42 (9.8%)",   "12", "480",   "58"],
            ["血液およびリンパ系障害",           "35 (8.2%)",   "8",  "420",   "38"],
            ["代謝および栄養障害",              "28 (6.5%)",   "4",  "310",   "18"],
            ["筋骨格系および結合組織障害",       "22 (5.1%)",   "1",  "245",   "6"],
            ["感染症および寄生虫症",            "18 (4.2%)",   "12", "215",   "58"],
            ["腎および尿路障害",                "12 (2.8%)",   "4",  "135",   "22"],
            ["その他",                          "68 (15.9%)",  "10", "780",   "42"],
        ])
    doc.add_paragraph()
    _fig(doc, _ae_bar_png(),
         "図2. 器官別大分類 (SOC) 別 有害事象発現頻度 (臨床試験における両投与群比較、模擬データ)")

    _h(doc, "8. 重篤有害事象詳細一覧", 1)
    _p(doc, "報告期間中に本剤との因果関係が否定できないと判断された重篤有害事象の主要症例を以下に示す。"
             "全症例について、症例安全性報告書 (CIOMS-I 形式) を PMDA へ提出済みである。")
    _fill_table(doc,
        ["症例ID", "SOC / PT", "発現時期", "重症度", "因果関係", "転帰"],
        [
            ["DBX-2025-0842", "胃腸障害 / 消化管出血",         "投与14日",  "重度",   "関連あり",     "回復"],
            ["DBX-2025-0917", "肝胆道系障害 / 肝機能異常",     "投与28日",  "中等度", "関連の可能性", "回復"],
            ["DBX-2025-1058", "血液系 / 好中球減少症",         "投与42日",  "中等度", "関連あり",     "回復"],
            ["DBX-2025-1112", "胃腸障害 / 腸閉塞",             "投与60日",  "重度",   "関連なし",     "後遺症あり"],
            ["DBX-2026-0053", "肝胆道系障害 / 肝細胞損傷",     "投与35日",  "重度",   "関連あり",     "回復"],
            ["DBX-2026-0187", "感染症 / 好中球減少性発熱",     "投与49日",  "重度",   "関連あり",     "回復"],
            ["DBX-2026-0231", "神経系 / 末梢神経障害",         "投与84日",  "中等度", "関連の可能性", "軽快"],
        ])
    doc.add_paragraph()

    _h(doc, "9. シグナル評価", 1)
    _h(doc, "9.1 肝毒性シグナルの評価", 2)
    _p(doc, "報告期間中に発現した肝胆道系有害事象 42 件について、詳細な因果関係評価を実施した。"
             "Hy's Law 基準 (ALT/AST > 3×ULN かつ総ビリルビン > 2×ULN) を満たす症例は 3 例発現した。"
             "全例で投与中止により速やかな回復が確認されており、本剤の既知のリスクの範囲内と判断される。")

    _h(doc, "9.2 消化管毒性シグナルの評価", 2)
    _p(doc, "消化管出血 (n=4) および腸閉塞 (n=2) の発現について、症例ベースで検討した。"
             "いずれも既知の腫瘍浸潤または既往の消化管疾患を有する患者で発現しており、"
             "本剤との直接的な因果関係は明確でないが、"
             "既知のリスクとして添付文書 4.4 項に注意喚起を継続する。")

    _h(doc, "9.3 未検出シグナル", 2)
    _p(doc, "報告期間中、以下のいずれの領域においても新たなシグナルは検出されなかった:"
             "(1) 心血管系有害事象、(2) 皮膚重症有害事象 (SJS/TEN 等)、(3) 免疫関連有害事象、"
             "(4) 二次発癌、(5) 生殖毒性関連事象。")

    _h(doc, "10. 特別な患者集団の安全性", 1)
    _h(doc, "10.1 高齢者", 2)
    _p(doc, "65 歳以上の高齢者集団 (n=1,240、全体の 38.3%) における有害事象発現頻度は 74.6% であり、"
             "全体集団の 71.8% と大きな差異は認められなかった。"
             "重篤 AE 発現率は高齢者で 18.2%、非高齢者で 12.4% と、"
             "やや高齢者で高い傾向が認められたが、統計学的有意差はなかった。")

    _h(doc, "10.2 腎機能障害患者", 2)
    _p(doc, "軽度〜中等度腎機能障害 (CrCl 30-89 mL/min) を有する患者では、"
             "通常用量での忍容性が確認されている。重度腎機能障害 (CrCl < 30 mL/min) を有する患者"
             "におけるデータは限定的であり、慎重投与を継続する。")

    _h(doc, "10.3 肝機能障害患者", 2)
    _p(doc, "Child-Pugh Class A の軽度肝機能障害患者では、通常用量での忍容性が確認されている。"
             "Class B の中等度肝機能障害患者では、100 mg 開始用量を推奨する。"
             "Class C の重度肝機能障害患者に対する使用は禁忌とする。")

    _h(doc, "11. 妊娠中および授乳中の安全性", 1)
    _p(doc, "本剤は生殖毒性を有することが動物試験で確認されており、妊娠中の使用は禁忌である。"
             "報告期間中に妊娠発現の報告が 1 例あり、患者・医療者への聴取結果、および転帰の追跡調査を実施中である。"
             "授乳中の安全性データはないため、授乳中の使用も禁忌としている。")

    _h(doc, "12. 累積安全性データ要約", 1)
    _p(doc, "承認以降の累積曝露 12,860 例における主要な安全性データを以下に要約する。"
             "全累積期間を通じて、既知のリスク以外に新たな安全性懸念は特定されていない。"
             "リスク・ベネフィットバランスは引き続き良好と判断される。")

    _h(doc, "13. 市販後調査 (PMS) 状況", 1)
    _p(doc, "国内市販後調査 (実施期間 3 年間、目標登録 3,000 例) は登録が順調に進捗しており、"
             "報告期間終了時点で 1,842 例が登録完了、1,240 例が観察期間を完了している。"
             "現時点で予期せぬ有害事象の発現は認められていない。中間解析は 2026 年下期に実施予定である。")

    _h(doc, "14. 結論および次回報告予定", 1)
    _p(doc, "本報告期間において、DBX-101 のリスク・ベネフィットバランスは引き続き良好と判断される。"
             "既知のリスク (消化器毒性、肝毒性、血液毒性) については、"
             "現行の添付文書記載および注意喚起により適切に管理されていると考えられる。"
             "現時点で規制上の追加措置を要する情報は認められない。")
    _p(doc, "次回 PSUR は、2026 年 10 月 1 日データロック、2026 年 12 月末までの提出予定である。"
             "以降 6 か月毎の定期的な安全性情報の集積および分析を継続する。")

    out = OUT_DIR / "02_ae_summary_ja.docx"
    doc.save(out)
    return out


# ---------------------------------------------------------------------------
# Doc 3 — CTD Module 2.7.3 Clinical Efficacy Summary (JA, ~10 pages)
# Structure: ICH M4E (R2)
# ---------------------------------------------------------------------------

def build_ctd_efficacy() -> Path:
    doc = Document()
    _add_header_footer(
        doc,
        header_text="CTD 第2.7.3項  臨床的有効性の概要  |  DBX-101  |  版 1.0 (2026年6月)",
        footer_text="Databricks Pharma KK  申請資料",
    )

    _h(doc, "第2.7.3項 臨床的有効性の概要", 0)
    _p(doc, "本項では、DBX-101 (開発コード) の臨床開発プログラムにおいて実施された主要な有効性試験の結果を要約する。"
             "本要約は、ICH M4E (R2) Common Technical Document ガイドラインに基づいて作成されている。")

    _h(doc, "1. 臨床開発プログラムの背景", 1)
    _p(doc, "DBX-101 の臨床開発は、2022 年 4 月に第 I 相試験 (DBX101-JP-001) が日本で開始された。"
             "第 I 相試験では、単回投与および反復投与時の安全性、忍容性、および薬物動態が評価された。"
             "その結果、100 mg および 200 mg 用量で忍容性が確認され、"
             "第 II 相試験の推奨用量として設定された。")
    _p(doc, "続いて 2023 年 3 月には、対象疾患を有する患者を対象とした"
             "無作為化二重盲検比較試験 (DBX101-JP-002) が開始された。"
             "同時に、対象疾患のサブタイプに対する単群非盲検試験 (DBX101-GLB-003) が"
             "国際共同試験として開始された。")

    _h(doc, "2. 個別試験の概要", 1)
    _fill_table(doc,
        ["試験番号", "相", "デザイン", "対象", "登録数", "主要評価項目", "実施期間"],
        [
            ["DBX101-JP-001",  "第I相",  "非盲検 用量漸増",       "進行固形癌",           "18",  "最大耐用量",     "2022-04 〜 2023-02"],
            ["DBX101-JP-002",  "第II相", "無作為化 二重盲検 比較",  "対象疾患",             "180", "PFS",           "2023-03 〜 2026-04"],
            ["DBX101-GLB-003", "第II相", "非盲検 単群",           "対象疾患サブタイプA",   "45",  "ORR",           "2023-05 〜 2026-05"],
            ["DBX101-JP-004",  "第I/II相", "併用療法 用量漸増",   "対象疾患 未治療",       "36",  "MTD + ORR",     "2024-06 〜 継続中"],
            ["DBX101-JPN-005", "特別調査", "観察研究",             "対象疾患 (実診療下)",    "1,842", "安全性・有効性", "2024-08 〜 継続中"],
        ])
    doc.add_paragraph()

    _h(doc, "3. 主要有効性結果 — DBX101-JP-002 試験", 1)
    _p(doc, "本申請の主要な有効性データは、無作為化二重盲検比較試験である DBX101-JP-002 試験から得られた。"
             "本試験では、標準治療に不応または不耐となった対象疾患患者 180 例を、"
             "DBX-101 100 mg 群、DBX-101 200 mg 群、または対照群 (標準治療) に 1:1:1 で無作為に割付けた。")
    _p(doc, "主要評価項目である独立中央判定 (BICR) に基づく無増悪生存期間 (PFS) の中央値は、"
             "DBX-101 高用量群で 10.4 か月 (95%CI: 8.6 - 12.1)、対照群で 6.5 か月 (95%CI: 5.2 - 7.8) であり、"
             "ハザード比 0.68 (95%CI: 0.55 - 0.84, 片側 p=0.0004) と統計学的に有意な延長が認められた。")

    _fill_table(doc,
        ["エンドポイント", "DBX-101 100 mg (n=60)", "DBX-101 200 mg (n=60)", "対照群 (n=60)", "HR (95%CI)", "p値"],
        [
            ["PFS 中央値 (月)",         "8.2 (6.9-9.8)",    "10.4 (8.6-12.1)", "6.5 (5.2-7.8)",   "0.68 (0.55-0.84)", "0.0004"],
            ["OS 中央値 (月)",          "18.3 (15.1-22.0)", "21.7 (18.2-25.9)", "16.1 (13.4-19.2)", "0.79 (0.61-1.02)", "0.068"],
            ["ORR n (%)",              "22 (36.7)",         "31 (51.7)",       "12 (20.0)",       "N/A",              "<0.001"],
            ["DCR n (%)",              "48 (80.0)",         "54 (90.0)",       "36 (60.0)",       "N/A",              "<0.001"],
            ["奏効期間 中央値 (月)",    "7.5 (5.8-9.6)",    "9.8 (7.2-12.4)",  "5.2 (3.6-7.1)",   "N/A",              "N/A"],
            ["治療期間 中央値 (月)",    "7.2 (5.4-9.1)",    "9.5 (7.0-12.3)",  "5.8 (4.1-7.6)",   "N/A",              "N/A"],
        ])
    doc.add_paragraph()

    _fig(doc, _km_curve_png(),
         "図1. 主要評価項目 (無増悪生存期間) の Kaplan-Meier 曲線 (投与群別、模擬データ)")
    _p(doc, "Kaplan-Meier 曲線に示されるとおり、DBX-101 200 mg 群では、対照群と比較して"
             "投与開始後全観察期間を通じて、無増悪生存確率が高く維持された。"
             "12 か月時点の PFS 率は、DBX-101 200 mg 群で 42%、100 mg 群で 32%、対照群で 18% であった。")

    _h(doc, "4. 部分集団解析", 1)
    _p(doc, "事前に規定された部分集団因子 (年齢、性別、ECOG PS、疾患重症度、前治療数等) について、"
             "主要評価項目のハザード比を評価した。"
             "全ての部分集団において、DBX-101 は対照群と比較して PFS を延長する方向性を示した。")
    _fig(doc, _forest_png(),
         "図2. 部分集団別 主要評価項目 (PFS) ハザード比のフォレストプロット (模擬データ)")
    _p(doc, "特に、前治療 1 レジメンの患者においてハザード比 0.55 (95%CI: 0.39 - 0.78) と、"
             "より強い治療効果が示唆された。ECOG PS 0 の患者、および軽症患者においても、"
             "全体集団を上回る効果が観察された。")

    _h(doc, "5. 副次評価項目 — 客観的奏効率 (ORR) および奏効期間", 1)
    _p(doc, "客観的奏効率 (ORR) は、DBX-101 200 mg 群で 51.7% (95%CI: 38.4 - 64.8)、"
             "100 mg 群で 36.7% (95%CI: 24.6 - 50.1)、対照群で 20.0% (95%CI: 10.8 - 32.3) であった。"
             "対照群に対する 200 mg 群の絶対差は 31.7% (95%CI: 15.2 - 46.2, p<0.001) と、"
             "統計学的に有意な改善が示された。")

    _fig(doc, _waterfall_png(),
         "図3. 腫瘍縮小効果 ウォーターフォールプロット (DBX-101 200 mg 群、模擬データ)")
    _p(doc, "ウォーターフォールプロットに示されるとおり、DBX-101 200 mg 群の 55 例中、"
             "PR 基準を満たす腫瘍縮小 (≥ 30% 縮小) が 28 例で確認された。"
             "また、進行 (PD) 基準を超える増大 (≥ 20% 増大) は 6 例に留まっており、"
             "対照群と比較して顕著に少なかった。")

    _h(doc, "6. 全生存期間 (OS) の考察", 1)
    _p(doc, "OS については、DBX-101 高用量群で対照群に比してハザード比 0.79 (95%CI: 0.61 - 1.02)"
             "と延長傾向が認められたが、本試験のデータカットオフ時点においては"
             "統計学的な有意差 (p=0.068) には至らなかった。"
             "OS イベント発生数は 105 件 (60% 発生) であり、追跡調査は継続中である。"
             "更新解析結果は、2026 年下期の申請資料補完で提出予定である。")

    _h(doc, "7. 用量反応関係の解析", 1)
    _p(doc, "DBX-101 100 mg 群と 200 mg 群の比較において、全ての主要な有効性エンドポイント"
             "(PFS、ORR、DCR、DOR) について、200 mg 群の方が数値的に良好な結果が観察された。"
             "この用量反応関係は、事前に規定された用量反応解析における Cochran-Armitage 検定でも"
             "統計学的有意 (傾向性検定 p<0.001) が示された。")
    _p(doc, "薬物動態解析においても、200 mg 群では 100 mg 群に比して"
             "定常状態 Ctrough が約 2.1 倍高値を示しており、"
             "血中濃度と有効性の相関 (E-R 関係) が支持される。"
             "以上より、推奨用量として 200 mg/日を提案する。")

    _h(doc, "8. サブタイプ試験 — DBX101-GLB-003 の結果", 1)
    _p(doc, "国際共同単群試験 DBX101-GLB-003 において、対象疾患のサブタイプ A 患者 45 例に対し、"
             "DBX-101 200 mg を投与した結果、主要評価項目である ORR は 62.2% (95%CI: 46.5 - 76.2) と、"
             "極めて高い奏効率が示された。DCR は 91.1% (95%CI: 78.8 - 97.5) であり、"
             "サブタイプ A における顕著な治療効果が確認された。")
    _fill_table(doc,
        ["エンドポイント", "DBX101-GLB-003 (サブタイプA)", "DBX101-JP-002 高用量群 (全対象)"],
        [
            ["ORR (%)",           "62.2 (46.5-76.2)",   "51.7 (38.4-64.8)"],
            ["DCR (%)",           "91.1 (78.8-97.5)",   "90.0 (79.5-96.2)"],
            ["PFS 中央値 (月)",    "13.8 (10.2-17.5)",  "10.4 (8.6-12.1)"],
            ["OS 中央値 (月)",     "未到達",             "21.7 (18.2-25.9)"],
        ])
    doc.add_paragraph()

    _h(doc, "9. 特別集団における有効性", 1)
    _p(doc, "高齢者 (65 歳以上) においても、非高齢者と同等の有効性が確認された。"
             "PFS 中央値は、65 歳以上で 10.1 か月、65 歳未満で 10.8 か月であり、"
             "ハザード比の交互作用検定は有意でなかった (p=0.42)。"
             "また、軽度腎機能障害を有する患者、および軽度肝機能障害を有する患者においても、"
             "全体集団と同等の有効性が示された。")

    _h(doc, "10. バイオマーカーサブ解析", 1)
    _p(doc, "本試験では、腫瘍組織および血漿検体を用いたバイオマーカー探索解析を実施した。"
             "標的経路の発現レベルが高値の集団 (n=52) では、"
             "PFS 中央値 13.2 か月 (95%CI: 10.4 - 16.5) と、"
             "発現レベル低値の集団 (n=68) の PFS 中央値 8.4 か月 (95%CI: 6.2 - 10.8) に比して"
             "顕著に長かった。今後、当該バイオマーカーの臨床導入について検討を進める。")

    _h(doc, "11. 追跡調査の状況", 1)
    _p(doc, "DBX101-JP-002 試験の追跡調査は、2026 年 4 月のデータカットオフ時点で継続中である。"
             "追跡完了予定は 2027 年 4 月であり、成熟した OS データおよび長期安全性データが得られる予定である。"
             "また、DBX101-GLB-003 試験の追跡も継続中であり、"
             "OS 中央値および長期奏効データは、2026 年下期に更新報告予定である。")

    _h(doc, "12. 有効性の総合考察", 1)
    _p(doc, "以上の複数試験からの結果を総合すると、DBX-101 は対象疾患を有する患者において、"
             "標準治療と比較して臨床的に意義のある無増悪生存期間の延長を示した。"
             "また、客観的奏効率、病勢制御率、奏効期間の全ての主要な有効性指標において、"
             "対照群に比して統計学的および臨床的に有意な改善が確認された。")
    _p(doc, "サブタイプ A 患者においては、単群試験の結果、極めて高い奏効率が示唆されており、"
             "本サブタイプに対する優先的な適応の可能性が示唆される。"
             "全生存期間については追跡調査中であるが、延長傾向が既に確認されており、"
             "成熟データによる統計学的有意性の確認が期待される。")

    _h(doc, "13. 結論", 1)
    _p(doc, "本項に示した有効性データと、第 2.7.4 項に示す安全性データを総合的に評価した結果、"
             "DBX-101 の臨床的有用性は、対象疾患を有する成人患者において良好であると結論する。"
             "本剤は、当該疾患領域における新規治療選択肢として、患者アウトカムの改善に寄与するものと期待される。")

    out = OUT_DIR / "03_ctd_efficacy_ja.docx"
    doc.save(out)
    return out


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    outs = [build_protocol_synopsis(), build_ae_summary(), build_ctd_efficacy()]
    for p in outs:
        size_kb = p.stat().st_size / 1024
        print(f"  wrote {p.name}  ({size_kb:.1f} KB)")
    print(f"\n{len(outs)} demo DOCX files ready in {OUT_DIR}")
